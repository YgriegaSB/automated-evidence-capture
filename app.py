import os
import csv
import time
import sys
from datetime import datetime
from pynput.keyboard import Key, Listener as KeyboardListener
from pynput.mouse import Listener as MouseListener
from PIL import ImageGrab, ImageDraw
from docx import Document
from docx.shared import Inches

running = True

if len(sys.argv) < 2:
    print("Por favor, proporcione el nombre del caso de prueba.")
    sys.exit(1)

test_case_name = sys.argv[1]

current_directory = os.path.dirname(os.path.abspath(__file__))
current_datetime = datetime.now().strftime('%y-%m-%d_%H-%M-%S')

evidence_capture_directory = os.path.join(current_directory, 'evidence-capture')
if not os.path.exists(evidence_capture_directory):
    os.makedirs(evidence_capture_directory)

evidence_folder = f'evidencia-{current_datetime}'
evidence_folder_path = os.path.join(evidence_capture_directory, evidence_folder)
os.makedirs(evidence_folder_path)

csv_file_path = os.path.join(evidence_folder_path, 'written_text.csv')
word_file_path = os.path.join(evidence_folder_path, f'{test_case_name}.docx')

step = 1
is_typing = False
typing_buffer = ""
last_key_time = time.time()

doc = Document()
doc.add_heading(test_case_name, 0)

def on_click(x, y, button, pressed):
    global step
    if pressed:
        screenshot = ImageGrab.grab()
        screenshot_path = os.path.join(evidence_folder_path, f'temp_step{step}.png')
        
        draw = ImageDraw.Draw(screenshot, 'RGBA')
        radius = 30
        red_transparent = (255, 0, 0, 100)
        draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=red_transparent)
        
        screenshot.save(screenshot_path)
        
        doc.add_heading(f'Paso {step}', level=1)
        doc.add_picture(screenshot_path, width=Inches(5))
        
        os.remove(screenshot_path)
        
        step += 1

def on_press(key):
    global is_typing, typing_buffer, last_key_time, running

    if key == Key.esc:
        running = False
        return False

    try:
        if hasattr(key, 'char'):
            typing_buffer += key.char
            is_typing = True
            last_key_time = time.time()
    except AttributeError:
        if key == Key.space:
            typing_buffer += ' '
        elif key == Key.enter:
            typing_buffer += '\n'

def on_release(key):
    global is_typing, typing_buffer, last_key_time

    if time.time() - last_key_time > 2 and is_typing:
        is_typing = False

        if typing_buffer.strip():
            save_text_to_csv(typing_buffer)
            take_screenshot()

        typing_buffer = ""

def save_text_to_csv(text):
    global step
    with open(csv_file_path, 'a', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow([step, text])
        step += 1

def take_screenshot():
    global step
    screenshot = ImageGrab.grab()
    screenshot_path = os.path.join(evidence_folder_path, f'temp_screenshot_{step}.png')
    screenshot.save(screenshot_path)
    
    doc.add_heading(f'Paso {step}', level=1)
    doc.add_picture(screenshot_path, width=Inches(5))
    
    os.remove(screenshot_path)
    
    step += 1

key_listener = KeyboardListener(on_press=on_press, on_release=on_release)
key_listener.start()

mouse_listener = MouseListener(on_click=on_click)
mouse_listener.start()

while running:
    time.sleep(0.1)

key_listener.stop()
mouse_listener.stop()

doc.add_heading('Observaciones', level=1)
doc.add_paragraph('Aquí puedes añadir cualquier observación relevante sobre el caso de prueba.')

doc.add_heading('Resultado de Certificación', level=1)
doc.add_paragraph('Resultado de la certificación: OK')

doc.save(word_file_path)
print(f"Informe generado: {word_file_path}")
